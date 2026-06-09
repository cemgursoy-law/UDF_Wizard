#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
udf_core.py
UYAP UDF <-> DOCX / PDF dönüştürme çekirdek fonksiyonları.
Tamamen offline çalışır. Harici servis kullanmaz.

UDF yapısı:
  - .udf = ZIP arşivi, içinde tek dosya: content.xml
  - content.xml:
      <template format_id="1.7">
        <content><![CDATA[ ...tüm düz metin, paragraflar \n ile ayrık... ]]></content>
        <properties><pageFormat .../></properties>
        <elements resolver="hvl-default">
          <paragraph Alignment="N" resolver="hvl-default">
            <content [bold="true"] size="12" family="Times New Roman"
                     startOffset="X" length="L" />
          </paragraph>
          ...
        </elements>
        <styles>
          <style name="default" .../>
          <style name="hvl-default" .../>
        </styles>
      </template>

Hizalama (Alignment) kodları: 0=sol, 1=orta, 2=sağ, 3=iki yana yasla
"""

import os
import re
import io
import base64
import zipfile
import xml.sax.saxutils as sax

# ---- isteğe bağlı kütüphaneler (uygulama bunlar yoksa da çekirdek kısmı çalışır) ----
try:
    import docx  # python-docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

try:
    import pdfplumber
    HAVE_PDFPLUMBER = True
except Exception:
    HAVE_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF - PDF içindeki görselleri çıkarmak için
    HAVE_FITZ = True
except Exception:
    HAVE_FITZ = False

try:
    from PIL import Image as PILImage  # görsellerin piksel boyutunu okumak için
    HAVE_PIL = True
except Exception:
    HAVE_PIL = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Image as RLImage
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAVE_REPORTLAB = True
except Exception:
    HAVE_REPORTLAB = False


ALIGN_DOCX_TO_UDF = {
    None: 0,
    0: 0,   # LEFT
    1: 1,   # CENTER
    2: 2,   # RIGHT
    3: 3,   # JUSTIFY
}
ALIGN_UDF_TO_DOCX = {0: 0, 1: 1, 2: 2, 3: 3}

IMAGE_PLACEHOLDER = "�"  # UYAP'ın gömülü görseller için içerik metninde bıraktığı yer tutucu karakter


def _image_pixel_size(data):
    """Görsel baytlarından (genişlik, yükseklik) piksel boyutunu döndürür; okunamazsa (None, None)."""
    if not HAVE_PIL or not data:
        return None, None
    try:
        with PILImage.open(io.BytesIO(data)) as im:
            return float(im.width), float(im.height)
    except Exception:
        return None, None


# =====================================================================
#  UDF YAZMA
# =====================================================================
def build_udf_content_xml(paragraphs):
    """
    paragraphs: list of dict.
      Metin   : {type:"text",  text, align (0-3), bold (bool)}
      Görsel  : {type:"image", data (bytes), width, height}
    Tek paragrafın tamamı tek biçimde varsayılır (basit ve sağlam model).
    """
    full = ""
    specs = []
    for para in paragraphs:
        start = len(full)
        if para.get("type") == "image":
            full += IMAGE_PLACEHOLDER + "\n"
            specs.append(("image", start, 1, para))
        else:
            text = para.get("text", "")
            full += text
            seg_len = len(text)
            full += "\n"  # her paragraf sonunda newline
            specs.append(("text", start, seg_len + 1, para))

    out = []
    out.append('<?xml version="1.0" encoding="UTF-8" ?>')
    out.append('<template format_id="1.7">')
    out.append('<content><![CDATA[' + full + ']]></content>')
    out.append('<properties><pageFormat mediaSizeName="1" '
               'leftMargin="56.7" rightMargin="56.699999999999974" '
               'topMargin="56.7" bottomMargin="56.699999999999974" '
               'paperOrientation="1" headerFOffset="20.0" footerFOffset="20.0" /></properties>')
    out.append('<elements resolver="hvl-default">')
    for kind, start, length, para in specs:
        if kind == "image":
            data = para.get("data") or b""
            w, h = para.get("width"), para.get("height")
            if w is None or h is None:
                w, h = _image_pixel_size(data)
            b64 = base64.encodebytes(data).decode("ascii")
            out.append('<paragraph>')
            out.append('<image imageData="%s" width="%s" height="%s" startOffset="%d" length="%d" />'
                       % (b64, w if w is not None else 0, h if h is not None else 0, start, length))
            out.append('</paragraph>')
        else:
            align, bold = para.get("align", 0), para.get("bold", False)
            cattr = ('bold="true" ' if bold else '') + 'size="12" family="Times New Roman" '
            out.append('<paragraph Alignment="%d" resolver="hvl-default">' % align)
            out.append('<content %sstartOffset="%d" length="%d" />' % (cattr, start, length))
            out.append('</paragraph>')
    out.append('</elements>')
    out.append('<styles>')
    out.append('<style name="default" description="Geçerli" italic="false" bold="false" '
               'FONT_ATTRIBUTE_KEY="javax.swing.plaf.FontUIResource[family=Tahoma,name=Tahoma,style=plain,size=11]" '
               'size="11" family="Tahoma" />')
    out.append('<style name="hvl-default" description="Gövde" size="12" family="Times New Roman" />')
    out.append('</styles>')
    out.append('</template>')
    return "\n".join(out)


def write_udf(paragraphs, out_path):
    xml_str = build_udf_content_xml(paragraphs)
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("content.xml", xml_str.encode("utf-8"))
    return out_path


# =====================================================================
#  UDF OKUMA
# =====================================================================
def read_udf(path):
    """UDF -> list of {type:"text", text, align, bold} / {type:"image", data, width, height}"""
    with zipfile.ZipFile(path, "r") as z:
        names = z.namelist()
        cname = "content.xml" if "content.xml" in names else names[0]
        raw = z.read(cname).decode("utf-8", errors="replace")

    # content (CDATA) çek
    m = re.search(r"<content>\s*<!\[CDATA\[(.*?)\]\]>\s*</content>", raw, re.S)
    full_text = m.group(1) if m else ""

    # paragraph + content offset bilgilerini çek
    paras = []
    for pm in re.finditer(r"<paragraph\b([^>]*)>(.*?)</paragraph>", raw, re.S):
        pattr = pm.group(1)
        inner = pm.group(2)

        im = re.search(r'<image\b([^>]*?)/?>', inner, re.S)
        if im:
            iattr = im.group(1)
            dm = re.search(r'imageData\s*=\s*"(.*?)"', iattr, re.S)
            wm = re.search(r'width\s*=\s*"([\d.]+)"', iattr)
            hm = re.search(r'height\s*=\s*"([\d.]+)"', iattr)
            try:
                data = base64.b64decode(re.sub(r"\s+", "", dm.group(1))) if dm else b""
            except Exception:
                data = b""
            paras.append({"type": "image", "data": data,
                          "width": float(wm.group(1)) if wm else None,
                          "height": float(hm.group(1)) if hm else None})
            continue

        am = re.search(r'Alignment="(\d+)"', pattr)
        align = int(am.group(1)) if am else 0
        # Bir paragrafta birden çok <content> parçası olabilir; hepsini birleştir.
        seg = ""
        bold_flags = []
        for cm in re.finditer(r'<content\b([^>]*?)/?>', inner):
            cattr = cm.group(1)
            sm = re.search(r'startOffset="(\d+)"', cattr)
            lm = re.search(r'length="(\d+)"', cattr)
            if sm and lm:
                s = int(sm.group(1)); l = int(lm.group(1))
                seg += full_text[s:s+l]
                bold_flags.append('bold="true"' in cattr)
        # paragraf, parçalarının çoğunluğu kalınsa kalın sayılır
        bold = bool(bold_flags) and (sum(bold_flags) * 2 >= len(bold_flags))
        seg = seg.rstrip("\n")
        paras.append({"type": "text", "text": seg, "align": align, "bold": bold})

    # elements yoksa düz metni satırlara böl
    if not paras and full_text:
        for line in full_text.split("\n"):
            paras.append({"type": "text", "text": line, "align": 0, "bold": False})
    return paras


# =====================================================================
#  DOCX <-> paragraphs
# =====================================================================
def _docx_paragraph_images(paragraph):
    """Bir paragraf içindeki gömülü (inline) görselleri {type, data, width, height} olarak döndürür."""
    images = []
    for blip in paragraph._p.findall(".//" + qn("a:blip")):
        rId = blip.get(qn("r:embed"))
        if not rId:
            continue
        try:
            data = paragraph.part.related_parts[rId].blob
        except Exception:
            continue
        w, h = _image_pixel_size(data)
        images.append({"type": "image", "data": data, "width": w, "height": h})
    return images


def docx_to_paragraphs(path):
    if not HAVE_DOCX:
        raise RuntimeError("python-docx yüklü değil")
    d = docx.Document(path)
    paras = []
    for p in d.paragraphs:
        images = _docx_paragraph_images(p)
        paras.extend(images)
        text = p.text
        if text.strip() or not images:
            align = ALIGN_DOCX_TO_UDF.get(int(p.alignment) if p.alignment is not None else None, 0)
            # paragraf kalın mı: tüm run'lar bold ise kalın say
            runs = [r for r in p.runs if r.text.strip()]
            bold = bool(runs) and all(r.bold for r in runs)
            paras.append({"type": "text", "text": text, "align": align, "bold": bold})
    return paras


def paragraphs_to_docx(paragraphs, out_path):
    if not HAVE_DOCX:
        raise RuntimeError("python-docx yüklü değil")
    d = docx.Document()
    # varsayılan font
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    align_map = {0: WD_ALIGN_PARAGRAPH.LEFT, 1: WD_ALIGN_PARAGRAPH.CENTER,
                 2: WD_ALIGN_PARAGRAPH.RIGHT, 3: WD_ALIGN_PARAGRAPH.JUSTIFY}
    for para in paragraphs:
        if para.get("type") == "image":
            data = para.get("data")
            if data:
                try:
                    d.add_picture(io.BytesIO(data))
                except Exception:
                    pass
            continue
        p = d.add_paragraph()
        p.alignment = align_map.get(para.get("align", 0), WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(para.get("text", ""))
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        if para.get("bold"):
            r.bold = True
    d.save(out_path)
    return out_path


# =====================================================================
#  PDF -> paragraphs  /  paragraphs -> PDF
# =====================================================================
def _pdf_page_images(fitz_doc, page_index):
    """Bir PDF sayfasındaki gömülü görselleri {type, data, width, height} olarak döndürür."""
    out = []
    page = fitz_doc[page_index]
    seen = set()
    for img in page.get_images(full=True):
        xref = img[0]
        if xref in seen:
            continue
        seen.add(xref)
        try:
            info = fitz_doc.extract_image(xref)
            data = info.get("image")
        except Exception:
            data = None
        if not data:
            continue
        out.append({"type": "image", "data": data,
                    "width": float(info.get("width") or 0),
                    "height": float(info.get("height") or 0)})
    return out


def pdf_to_paragraphs(path):
    if not HAVE_PDFPLUMBER:
        raise RuntimeError("pdfplumber yüklü değil")
    fitz_doc = None
    if HAVE_FITZ:
        try:
            fitz_doc = fitz.open(path)
        except Exception:
            fitz_doc = None
    try:
        paras = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                txt = page.extract_text() or ""
                for line in txt.split("\n"):
                    paras.append({"type": "text", "text": line, "align": 0, "bold": False})
                if fitz_doc is not None and i < fitz_doc.page_count:
                    paras.extend(_pdf_page_images(fitz_doc, i))
        return paras
    finally:
        if fitz_doc is not None:
            fitz_doc.close()


def _register_font():
    """Sistemde Türkçe destekli bir TTF bulup reportlab'a kaydet."""
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/Library/Fonts/Times New Roman.ttf",
    ]
    for c in candidates:
        if os.path.exists(c):
            try:
                pdfmetrics.registerFont(TTFont("DocFont", c))
                return "DocFont"
            except Exception:
                continue
    return "Helvetica"  # son çare (Türkçe karakterde sorun olabilir)


def paragraphs_to_pdf(paragraphs, out_path):
    if not HAVE_REPORTLAB:
        raise RuntimeError("reportlab yüklü değil")
    fontname = _register_font()
    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=20*mm, rightMargin=15*mm,
                            topMargin=20*mm, bottomMargin=15*mm)
    align_map = {0: TA_LEFT, 1: TA_CENTER, 2: TA_RIGHT, 3: TA_JUSTIFY}
    story = []
    for para in paragraphs:
        if para.get("type") == "image":
            data = para.get("data")
            if not data:
                continue
            w, h = para.get("width") or 0, para.get("height") or 0
            if not w or not h:
                w, h = _image_pixel_size(data)
                w, h = w or 0, h or 0
            if w and h and w > doc.width:
                scale = doc.width / w
                w, h = w * scale, h * scale
            try:
                if w and h:
                    story.append(RLImage(io.BytesIO(data), width=w, height=h))
                else:
                    story.append(RLImage(io.BytesIO(data)))
            except Exception:
                continue
            story.append(Spacer(1, 4))
            continue
        text = sax.escape(para.get("text", "")) or " "
        st = ParagraphStyle(
            "x", fontName=fontname, fontSize=12, leading=16,
            alignment=align_map.get(para.get("align", 0), TA_LEFT),
        )
        if para.get("bold"):
            text = "<b>%s</b>" % text
        story.append(RLParagraph(text.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;"), st))
        story.append(Spacer(1, 4))
    doc.build(story)
    return out_path


# =====================================================================
#  ÜST SEVİYE DÖNÜŞTÜRME
# =====================================================================
def convert(in_path, out_path):
    ie = os.path.splitext(in_path)[1].lower()
    oe = os.path.splitext(out_path)[1].lower()

    # önce kaynağı paragraflara çevir
    if ie == ".udf":
        paras = read_udf(in_path)
    elif ie == ".docx":
        paras = docx_to_paragraphs(in_path)
    elif ie == ".pdf":
        paras = pdf_to_paragraphs(in_path)
    else:
        raise ValueError("Desteklenmeyen girdi türü: " + ie)

    # sonra hedefe yaz
    if oe == ".udf":
        return write_udf(paras, out_path)
    elif oe == ".docx":
        return paragraphs_to_docx(paras, out_path)
    elif oe == ".pdf":
        return paragraphs_to_pdf(paras, out_path)
    else:
        raise ValueError("Desteklenmeyen çıktı türü: " + oe)


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Kullanım: python udf_core.py <girdi> <çıktı>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
    print("Tamam:", sys.argv[2])
